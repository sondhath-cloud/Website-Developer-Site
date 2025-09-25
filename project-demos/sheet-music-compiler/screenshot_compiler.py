import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
from datetime import datetime
from PIL import Image, ImageTk
from PIL.ExifTags import TAGS
import docx
from docx.shared import Inches
import glob
import numpy as np
from collections import Counter

def is_sheet_music(image_path):
    """
    Detect if an image is likely to be sheet music based on visual characteristics.
    Returns True if the image appears to be sheet music, False otherwise.
    """
    try:
        # Open and convert image to grayscale for analysis
        image = Image.open(image_path)
        
        # Convert to grayscale for analysis
        gray_image = image.convert('L')
        
        # Resize to standard size for analysis (faster processing)
        gray_image = gray_image.resize((400, 600), Image.Resampling.LANCZOS)
        
        # Convert to numpy array
        img_array = np.array(gray_image)
        
        # Calculate some basic characteristics
        height, width = img_array.shape
        
        # Check aspect ratio (sheet music is typically taller than wide)
        aspect_ratio = height / width
        if aspect_ratio < 1.2:  # Too wide, probably not sheet music
            return False
        
        # Analyze pixel distribution
        # Sheet music typically has:
        # 1. High contrast (black notes/staff lines on white background)
        # 2. Many horizontal lines (staff lines)
        # 3. Small black elements (notes, symbols)
        
        # Calculate contrast
        pixel_std = np.std(img_array)
        if pixel_std < 30:  # Low contrast, probably not sheet music
            return False
        
        # Look for horizontal lines (staff lines)
        # Apply horizontal edge detection
        horizontal_edges = np.abs(np.diff(img_array, axis=1))
        horizontal_strength = np.mean(horizontal_edges)
        
        # Look for vertical elements (note stems, bar lines)
        vertical_edges = np.abs(np.diff(img_array, axis=0))
        vertical_strength = np.mean(vertical_edges)
        
        # Sheet music should have more horizontal structure (staff lines)
        if horizontal_strength > vertical_strength * 1.2:
            return True
        
        # Additional check: look for repetitive horizontal patterns
        # Sample horizontal lines and check for staff-like patterns
        sample_lines = []
        for i in range(0, height, height // 10):
            line = img_array[i, :]
            sample_lines.append(line)
        
        # Check if we have the typical 5-line staff pattern
        staff_pattern_score = 0
        for line in sample_lines:
            # Look for lines with alternating light/dark patterns
            line_variation = np.std(line)
            if line_variation > 20:  # Significant variation suggests staff lines
                staff_pattern_score += 1
        
        if staff_pattern_score >= 3:  # Multiple lines with staff-like patterns
            return True
        
        # If we get here, it's probably not sheet music
        return False
        
    except Exception as e:
        # If we can't analyze the image, assume it might be sheet music
        # (better to include than exclude)
        return True

class GridPreviewWindow:
    def __init__(self, parent, image_data, callback):
        self.parent = parent
        self.image_data = image_data.copy()
        self.callback = callback
        self.selected_images = {}  # Track which images are selected
        
        self.preview_window = tk.Toplevel(parent)
        self.preview_window.title("Select Sheet Music Pages")
        self.preview_window.geometry("1000x700")
        self.preview_window.configure(bg='#2b2b2b')
        
        self.setup_grid_ui()
        self.display_thumbnails()
        
    def setup_grid_ui(self):
        # Title
        title_label = ttk.Label(self.preview_window, text="Select Screenshots to Compile", 
                               font=('Arial', 14, 'bold'))
        title_label.pack(pady=10)
        
        # Info label
        self.info_label = ttk.Label(self.preview_window, text=f"Found {len(self.image_data)} screenshots - Select which ones to include")
        self.info_label.pack(pady=5)
        
        # Selection controls
        selection_frame = tk.Frame(self.preview_window, bg='#2b2b2b')
        selection_frame.pack(pady=5)
        
        ttk.Button(selection_frame, text="Select All", 
                  command=self.select_all).pack(side='left', padx=5)
        ttk.Button(selection_frame, text="Select None", 
                  command=self.select_none).pack(side='left', padx=5)
        
        # Create scrollable frame for thumbnails
        canvas = tk.Canvas(self.preview_window, bg='#2b2b2b')
        scrollbar = ttk.Scrollbar(self.preview_window, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg='#2b2b2b')
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True, padx=10, pady=10)
        scrollbar.pack(side="right", fill="y")
        
        self.thumbnail_frame = scrollable_frame
        self.canvas = canvas
        
        # Bind mousewheel to canvas
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        # Action buttons
        button_frame = tk.Frame(self.preview_window, bg='#2b2b2b')
        button_frame.pack(pady=10)
        
        self.compile_button = ttk.Button(button_frame, text="Compile Selected Pages (0)", 
                                        command=self.compile_selected, state='disabled')
        self.compile_button.pack(side='left', padx=10)
        ttk.Button(button_frame, text="Cancel", 
                  command=self.preview_window.destroy).pack(side='left', padx=10)
        
    def display_thumbnails(self):
        """Display thumbnails of all screenshots with checkboxes"""
        # Clear existing thumbnails
        for widget in self.thumbnail_frame.winfo_children():
            widget.destroy()
        
        # Create grid of thumbnails
        cols = 3  # 3 columns
        for i, (date_time, image_path) in enumerate(self.image_data):
            row = i // cols
            col = i % cols
            
            # Create thumbnail frame
            thumb_frame = tk.Frame(self.thumbnail_frame, bg='#404040', relief='raised', bd=2)
            thumb_frame.grid(row=row, column=col, padx=5, pady=5, sticky='nsew')
            
            try:
                # Load and resize image for thumbnail
                image = Image.open(image_path)
                
                # Calculate thumbnail size
                thumb_size = 180
                img_width, img_height = image.size
                aspect_ratio = img_width / img_height
                
                if aspect_ratio > 1:
                    thumb_width = thumb_size
                    thumb_height = int(thumb_size / aspect_ratio)
                else:
                    thumb_height = thumb_size
                    thumb_width = int(thumb_size * aspect_ratio)
                
                # Add checkbox first
                var = tk.BooleanVar()
                checkbox = ttk.Checkbutton(thumb_frame, text="Include", variable=var,
                                         command=lambda idx=i, v=var: self.update_selection(idx, v))
                checkbox.pack(pady=2)
                
                # Store the checkbox variable
                self.selected_images[i] = var
                
                # Resize image
                thumb_image = image.resize((thumb_width, thumb_height), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(thumb_image)
                
                # Create image label
                img_label = tk.Label(thumb_frame, image=photo, bg='#404040')
                img_label.image = photo  # Keep reference
                img_label.pack(pady=5)
                
                # Add page info
                info_text = f"Screenshot {i+1}\n{date_time.strftime('%m/%d %H:%M')}"
                info_label = tk.Label(thumb_frame, text=info_text, bg='#404040', 
                                    fg='white', font=('Arial', 8))
                info_label.pack(pady=2)
                
                # Add click handler for full-size preview
                img_label.bind('<Button-1>', lambda e, idx=i: self.show_full_size_preview(idx))
                
            except Exception as e:
                # Show error thumbnail
                error_label = tk.Label(thumb_frame, text=f"Error loading\n{os.path.basename(image_path)}", 
                                     bg='#404040', fg='red')
                error_label.pack(pady=20)
        
        # Configure grid weights
        for i in range(cols):
            self.thumbnail_frame.columnconfigure(i, weight=1)
    
    def show_full_size_preview(self, index):
        """Show full-size preview of selected image"""
        date_time, image_path = self.image_data[index]
        
        # Create preview window
        preview = tk.Toplevel(self.preview_window)
        preview.title(f"Page {index + 1} - {date_time.strftime('%Y-%m-%d %H:%M:%S')}")
        preview.geometry("800x600")
        
        try:
            image = Image.open(image_path)
            
            # Calculate display size
            canvas_width = 750
            canvas_height = 550
            
            img_width, img_height = image.size
            aspect_ratio = img_width / img_height
            
            if aspect_ratio > canvas_width / canvas_height:
                display_width = canvas_width
                display_height = int(canvas_width / aspect_ratio)
            else:
                display_height = canvas_height
                display_width = int(canvas_height * aspect_ratio)
            
            # Resize image
            display_image = image.resize((display_width, display_height), Image.Resampling.LANCZOS)
            photo = ImageTk.PhotoImage(display_image)
            
            # Create canvas with scrollbars
            canvas = tk.Canvas(preview, width=canvas_width, height=canvas_height, bg='white')
            canvas.pack(padx=10, pady=10)
            
            canvas.create_image(canvas_width//2, canvas_height//2, image=photo)
            canvas.image = photo  # Keep reference
            
        except Exception as e:
            error_label = tk.Label(preview, text=f"Error loading image:\n{str(e)}", 
                                 fg='red', font=('Arial', 12))
            error_label.pack(pady=50)
    
    def update_selection(self, index, var):
        """Update the selection count and button state"""
        selected_count = sum(1 for v in self.selected_images.values() if v.get())
        self.compile_button.config(text=f"Compile Selected Pages ({selected_count})")
        
        if selected_count > 0:
            self.compile_button.config(state='normal')
        else:
            self.compile_button.config(state='disabled')
    
    def select_all(self):
        """Select all images"""
        for var in self.selected_images.values():
            var.set(True)
        self.update_selection(0, None)
    
    def select_none(self):
        """Deselect all images"""
        for var in self.selected_images.values():
            var.set(False)
        self.update_selection(0, None)
    
    def compile_selected(self):
        """Compile only the selected images in chronological order"""
        selected_data = []
        for i, var in self.selected_images.items():
            if var.get():
                selected_data.append(self.image_data[i])
        
        if selected_data:
            self.callback(selected_data)
            self.preview_window.destroy()
        else:
            messagebox.showwarning("No Images", "No images selected for compilation")

class PreviewWindow:
    def __init__(self, parent, image_data, callback):
        self.parent = parent
        self.image_data = image_data.copy()
        self.original_data = image_data.copy()
        self.callback = callback
        
        self.preview_window = tk.Toplevel(parent)
        self.preview_window.title("Review Sheet Music Images")
        self.preview_window.geometry("800x600")
        self.preview_window.configure(bg='#2b2b2b')
        
        self.current_index = 0
        self.setup_preview_ui()
        self.show_current_image()
        
    def setup_preview_ui(self):
        # Title
        title_label = ttk.Label(self.preview_window, text="Review Sheet Music Images", 
                               font=('Arial', 14, 'bold'))
        title_label.pack(pady=10)
        
        # Image info frame
        info_frame = tk.Frame(self.preview_window, bg='#2b2b2b')
        info_frame.pack(pady=5)
        
        self.info_label = ttk.Label(info_frame, text="")
        self.info_label.pack()
        
        # Main image display frame
        image_frame = tk.Frame(self.preview_window, bg='#2b2b2b')
        image_frame.pack(pady=10, padx=20, fill='both', expand=True)
        
        # Image label with scrollbars
        canvas_frame = tk.Frame(image_frame, bg='#2b2b2b')
        canvas_frame.pack(fill='both', expand=True)
        
        self.canvas = tk.Canvas(canvas_frame, bg='white', width=600, height=400)
        self.canvas.pack(side='left', fill='both', expand=True)
        
        # Scrollbars
        v_scrollbar = ttk.Scrollbar(canvas_frame, orient='vertical', command=self.canvas.yview)
        v_scrollbar.pack(side='right', fill='y')
        self.canvas.configure(yscrollcommand=v_scrollbar.set)
        
        h_scrollbar = ttk.Scrollbar(image_frame, orient='horizontal', command=self.canvas.xview)
        h_scrollbar.pack(side='bottom', fill='x')
        self.canvas.configure(xscrollcommand=h_scrollbar.set)
        
        # Navigation controls
        nav_frame = tk.Frame(self.preview_window, bg='#2b2b2b')
        nav_frame.pack(pady=10)
        
        ttk.Button(nav_frame, text="◀ Previous", command=self.prev_image).pack(side='left', padx=5)
        ttk.Button(nav_frame, text="Next ▶", command=self.next_image).pack(side='left', padx=5)
        
        # Control buttons frame
        control_frame = tk.Frame(self.preview_window, bg='#2b2b2b')
        control_frame.pack(pady=10)
        
        ttk.Button(control_frame, text="Move Up", command=self.move_up).pack(side='left', padx=5)
        ttk.Button(control_frame, text="Move Down", command=self.move_down).pack(side='left', padx=5)
        ttk.Button(control_frame, text="Remove from List", command=self.remove_current).pack(side='left', padx=5)
        
        # Action buttons
        action_frame = tk.Frame(self.preview_window, bg='#2b2b2b')
        action_frame.pack(pady=20)
        
        ttk.Button(action_frame, text="Compile Selected Images", 
                  command=self.compile_selected).pack(side='left', padx=10)
        ttk.Button(action_frame, text="Cancel", 
                  command=self.preview_window.destroy).pack(side='left', padx=10)
        
    def show_current_image(self):
        if not self.image_data:
            self.info_label.config(text="No images to display")
            return
            
        date_time, image_path = self.image_data[self.current_index]
        self.info_label.config(text=f"Image {self.current_index + 1} of {len(self.image_data)} - {date_time.strftime('%Y-%m-%d %H:%M:%S')}")
        
        try:
            # Load and resize image for display
            image = Image.open(image_path)
            
            # Calculate display size while maintaining aspect ratio
            canvas_width = 600
            canvas_height = 400
            
            img_width, img_height = image.size
            aspect_ratio = img_width / img_height
            
            if aspect_ratio > canvas_width / canvas_height:
                # Image is wider
                display_width = canvas_width
                display_height = int(canvas_width / aspect_ratio)
            else:
                # Image is taller
                display_height = canvas_height
                display_width = int(canvas_height * aspect_ratio)
            
            # Resize image
            display_image = image.resize((display_width, display_height), Image.Resampling.LANCZOS)
            
            # Convert to PhotoImage
            photo = ImageTk.PhotoImage(display_image)
            
            # Clear canvas and add image
            self.canvas.delete("all")
            self.canvas.create_image(canvas_width//2, canvas_height//2, image=photo)
            
            # Update scroll region
            self.canvas.configure(scrollregion=self.canvas.bbox("all"))
            
            # Keep reference to prevent garbage collection
            self.canvas.image = photo
            
        except Exception as e:
            self.canvas.delete("all")
            self.canvas.create_text(300, 200, text=f"Error loading image:\n{str(e)}", 
                                   fill="red", font=('Arial', 12))
    
    def prev_image(self):
        if self.current_index > 0:
            self.current_index -= 1
            self.show_current_image()
    
    def next_image(self):
        if self.current_index < len(self.image_data) - 1:
            self.current_index += 1
            self.show_current_image()
    
    def move_up(self):
        if self.current_index > 0:
            # Swap with previous item
            self.image_data[self.current_index], self.image_data[self.current_index - 1] = \
                self.image_data[self.current_index - 1], self.image_data[self.current_index]
            self.current_index -= 1
            self.show_current_image()
    
    def move_down(self):
        if self.current_index < len(self.image_data) - 1:
            # Swap with next item
            self.image_data[self.current_index], self.image_data[self.current_index + 1] = \
                self.image_data[self.current_index + 1], self.image_data[self.current_index]
            self.current_index += 1
            self.show_current_image()
    
    def remove_current(self):
        if len(self.image_data) > 1:
            del self.image_data[self.current_index]
            if self.current_index >= len(self.image_data):
                self.current_index = len(self.image_data) - 1
            self.show_current_image()
        else:
            messagebox.showwarning("Cannot Remove", "Cannot remove the last image")
    
    def compile_selected(self):
        if self.image_data:
            self.callback(self.image_data)
            self.preview_window.destroy()
        else:
            messagebox.showwarning("No Images", "No images selected for compilation")

class ScreenshotCompiler:
    def __init__(self, root):
        self.root = root
        self.root.title("Sheet Music Compiler")
        self.root.geometry("500x300")
        self.root.configure(bg='#2b2b2b')
        
        # Configure style for dark theme
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('TButton', background='#404040', foreground='white')
        style.configure('TLabel', background='#2b2b2b', foreground='white')
        style.configure('TCheckbutton', background='#2b2b2b', foreground='white')
        
        # Fixed folder paths
        self.source_folder = "/Users/sondrahealyhathaway/Desktop/Screenshots"
        self.output_folder = "/Users/sondrahealyhathaway/Documents/Personal/Sheet Music"
        
        self.setup_ui()
        
    def setup_ui(self):
        # Main title
        title_label = ttk.Label(self.root, text="Sheet Music Compiler", 
                               font=('Arial', 16, 'bold'))
        title_label.pack(pady=20)
        
        # Instructions
        instructions = ttk.Label(self.root, 
                               text="Compile sheet music screenshots into a Word document",
                               font=('Arial', 10))
        instructions.pack(pady=10)
        
        # Folder info frame
        folder_frame = tk.Frame(self.root, bg='#2b2b2b')
        folder_frame.pack(pady=20)
        
        # Source folder info
        source_label = ttk.Label(folder_frame, text="Source: Desktop/Screenshots")
        source_label.pack(pady=5)
        
        # Output folder info
        output_label = ttk.Label(folder_frame, text="Output: Documents/Personal/Sheet Music")
        output_label.pack(pady=5)
        
        # Compile button
        self.compile_button = ttk.Button(self.root, text="Load & Review Screenshots", 
                                       command=self.load_screenshots)
        self.compile_button.pack(pady=20)
        
        # Progress bar
        self.progress = ttk.Progressbar(self.root, mode='indeterminate')
        self.progress.pack(pady=10, padx=50, fill='x')
        
        # Status label
        self.status_label = ttk.Label(self.root, text="Ready to load screenshots")
        self.status_label.pack(pady=10)
        
    def load_screenshots(self):
        """Load screenshots from the fixed source folder"""
        # Start progress bar
        self.progress.start()
        self.status_label.config(text="Loading screenshots...")
        self.root.update()
        
        try:
            # Check if source folder exists
            if not os.path.exists(self.source_folder):
                messagebox.showerror("Error", f"Source folder not found:\n{self.source_folder}")
                self.progress.stop()
                self.status_label.config(text="Ready to load screenshots")
                return
            
            # Find all image files
            image_files = self.find_image_files(self.source_folder)
            
            if not image_files:
                messagebox.showwarning("No Images", "No image files found in the screenshots folder")
                self.progress.stop()
                self.status_label.config(text="Ready to load screenshots")
                return
            
            self.status_label.config(text=f"Found {len(image_files)} images. Sorting by date...")
            self.root.update()
            
            # Sort images by date/time
            image_data = []
            for image_path in image_files:
                date_time = self.get_image_date(image_path)
                image_data.append((date_time, image_path))
            
            # Sort by date/time (chronological order)
            image_data.sort(key=lambda x: x[0])
            
            # Stop progress bar
            self.progress.stop()
            self.status_label.config(text="Opening selection window...")
            self.root.update()
            
            # Open grid preview window
            GridPreviewWindow(self.root, image_data, self.compile_from_preview)
            
            self.status_label.config(text="Ready to load screenshots")
                
        except Exception as e:
            self.progress.stop()
            self.status_label.config(text="Error occurred")
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
            
    def get_image_date(self, image_path):
        """Extract date/time from image metadata or file modification time"""
        try:
            # Try to get EXIF data first
            image = Image.open(image_path)
            exif_data = image._getexif()
            
            if exif_data:
                for tag_id, value in exif_data.items():
                    tag = TAGS.get(tag_id, tag_id)
                    if tag == 'DateTime':
                        return datetime.strptime(value, '%Y:%m:%d %H:%M:%S')
                    elif tag == 'DateTimeOriginal':
                        return datetime.strptime(value, '%Y:%m:%d %H:%M:%S')
            
            # Fall back to file modification time
            timestamp = os.path.getmtime(image_path)
            return datetime.fromtimestamp(timestamp)
            
        except Exception as e:
            # If all else fails, use file modification time
            timestamp = os.path.getmtime(image_path)
            return datetime.fromtimestamp(timestamp)
    
    def find_image_files(self, folder):
        """Find all image files in the specified folder"""
        image_extensions = ['*.jpg', '*.jpeg', '*.png', '*.bmp', '*.tiff', '*.gif']
        image_files = []
        
        for extension in image_extensions:
            image_files.extend(glob.glob(os.path.join(folder, extension)))
            image_files.extend(glob.glob(os.path.join(folder, extension.upper())))
        
        return image_files
    
    
    def compile_from_preview(self, image_data):
        """Compile the selected images into a Word document"""
        if not image_data:
            messagebox.showwarning("No Images", "No images selected for compilation")
            return
            
        # Create output folder if it doesn't exist
        os.makedirs(self.output_folder, exist_ok=True)
            
        # Start progress bar
        self.progress.start()
        self.status_label.config(text="Creating Word document...")
        self.root.update()
        
        try:
            # Create Word document
            doc = docx.Document()
            doc.add_heading('Sheet Music Compilation', 0)
            doc.add_paragraph(f'Compiled on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Total pages: {len(image_data)}')
            doc.add_paragraph('')
            
            # Add images to document in chronological order
            for i, (date_time, image_path) in enumerate(image_data):
                self.status_label.config(text=f"Adding page {i+1} of {len(image_data)}...")
                self.root.update()
                
                try:
                    # Add page number
                    doc.add_heading(f'Page {i+1}', level=2)
                    
                    # Add the image
                    doc.add_picture(image_path, width=Inches(7.5))  # Larger for sheet music
                    
                    # Add some space
                    doc.add_paragraph('')
                    
                except Exception as e:
                    doc.add_paragraph(f'Error loading page: {os.path.basename(image_path)} - {str(e)}')
            
            # Save document
            output_filename = f"Sheet_Music_Compilation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(self.output_folder, output_filename)
            
            self.status_label.config(text="Saving document...")
            self.root.update()
            
            doc.save(output_path)
            
            # Stop progress bar
            self.progress.stop()
            self.status_label.config(text="Compilation complete!")
            
            # Show success message
            result = messagebox.askyesno("Success", 
                                       f"Sheet music document created successfully!\n\n"
                                       f"Saved as: {output_filename}\n"
                                       f"Location: {self.output_folder}\n\n"
                                       f"Would you like to open the folder?")
            
            if result:
                os.system(f'open "{self.output_folder}"')
                
        except Exception as e:
            self.progress.stop()
            self.status_label.config(text="Error occurred")
            messagebox.showerror("Error", f"An error occurred: {str(e)}")

def main():
    root = tk.Tk()
    app = ScreenshotCompiler(root)
    root.mainloop()

if __name__ == "__main__":
    main()
